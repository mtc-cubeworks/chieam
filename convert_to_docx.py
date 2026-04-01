#!/usr/bin/env python3
"""
Convert EAM-CHI USER_MANUAL.md and TEST_GUIDE.md to beautified DOCX files.
Features: styled tables, color-coded workflow diagrams, professional formatting.
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── Color Palette ──────────────────────────────────────────────────
BRAND_PRIMARY = RGBColor(0x1A, 0x56, 0xDB)   # Deep blue
BRAND_SECONDARY = RGBColor(0x0E, 0x7C, 0x86)  # Teal
BRAND_ACCENT = RGBColor(0xF5, 0x9E, 0x0B)     # Amber
BRAND_DARK = RGBColor(0x1F, 0x29, 0x37)       # Dark gray
BRAND_LIGHT = RGBColor(0xF8, 0xFA, 0xFC)      # Off-white
TABLE_HEADER_BG = "1A56DB"
TABLE_HEADER_BG_ALT = "0E7C86"
TABLE_ROW_EVEN = "F0F5FF"
TABLE_ROW_ODD = "FFFFFF"
TABLE_BORDER = "CBD5E1"
WORKFLOW_COLORS = {
    "Draft": ("94A3B8", "334155"),
    "Acquired": ("60A5FA", "1E3A5F"),
    "Inspected": ("A78BFA", "4C1D95"),
    "Active": ("34D399", "065F46"),
    "Inactive": ("F87171", "7F1D1D"),
    "Under Maintenance": ("FBBF24", "78350F"),
    "Under Repair": ("FB923C", "7C2D12"),
    "Decommissioned": ("6B7280", "1F2937"),
    "Pending Approval": ("FBBF24", "78350F"),
    "Approved": ("3B82F6", "1E3A5F"),
    "Release": ("8B5CF6", "4C1D95"),
    "Completed": ("22C55E", "065F46"),
    "Requested": ("60A5FA", "1E3A5F"),
    "In Progress": ("F59E0B", "78350F"),
    "Closed": ("10B981", "065F46"),
    "Awaiting Resources": ("94A3B8", "334155"),
    "Ready": ("3B82F6", "1E3A5F"),
    "On Hold": ("EF4444", "7F1D1D"),
    "Expired": ("6B7280", "1F2937"),
    "Cancelled": ("EF4444", "7F1D1D"),
    "Open": ("3B82F6", "1E3A5F"),
    "Rejected": ("EF4444", "7F1D1D"),
    "Warning": ("FBBF24", "78350F"),
    "Critical": ("EF4444", "7F1D1D"),
    "Resolved": ("22C55E", "065F46"),
    "Normal": ("10B981", "065F46"),
}


def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    tcPr.append(shading_elm)


def set_cell_borders(cell, color="CBD5E1", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:start w:w="{left}" w:type="dxa"/>'
        f'  <w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(margins)


def setup_document(doc, title_text, subtitle_text):
    """Set up document styles and properties."""
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Style: Normal
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = BRAND_DARK
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15

    # Style: Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(22)
    h1.font.color.rgb = BRAND_PRIMARY
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    # Style: Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(16)
    h2.font.color.rgb = BRAND_SECONDARY
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    # Style: Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(13)
    h3.font.color.rgb = BRAND_DARK
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(4)

    # Style: Heading 4
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Calibri'
    h4.font.size = Pt(11.5)
    h4.font.color.rgb = RGBColor(0x37, 0x47, 0x51)
    h4.font.bold = True
    h4.font.italic = True
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(4)

    # ── Cover Page ──
    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title_text)
    run.font.size = Pt(36)
    run.font.color.rgb = BRAND_PRIMARY
    run.font.bold = True
    run.font.name = 'Calibri'

    # Subtitle
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run(subtitle_text)
    run.font.size = Pt(16)
    run.font.color.rgb = BRAND_SECONDARY
    run.font.name = 'Calibri'

    # Info box
    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run("Enterprise Asset Management — CHI")
    run.font.size = Pt(14)
    run.font.color.rgb = BRAND_DARK
    run.font.name = 'Calibri'

    ver_para = doc.add_paragraph()
    ver_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver_para.add_run("Version 1.0  |  March 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    run.font.name = 'Calibri'

    # Decorative line
    doc.add_paragraph()
    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line_para.add_run("━" * 50)
    run.font.color.rgb = BRAND_PRIMARY
    run.font.size = Pt(12)

    # Page break after cover
    doc.add_page_break()

    return doc


def add_styled_table(doc, headers, rows, header_color=TABLE_HEADER_BG, compact=False):
    """Create a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Style header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        run.font.bold = True
        run.font.size = Pt(9.5) if compact else Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'
        set_cell_shading(cell, header_color)
        set_cell_borders(cell, header_color, "6")
        set_cell_margins(cell, 50, 50, 80, 80)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for row_idx, row_data in enumerate(rows):
        bg_color = TABLE_ROW_EVEN if row_idx % 2 == 0 else TABLE_ROW_ODD
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            text = str(cell_text).strip()

            # Check for icons / special formatting
            run = p.add_run(text)
            run.font.size = Pt(9) if compact else Pt(9.5)
            run.font.name = 'Calibri'
            run.font.color.rgb = BRAND_DARK

            # Bold first column for key-value tables
            if col_idx == 0 and len(headers) == 2:
                run.font.bold = True

            set_cell_shading(cell, bg_color)
            set_cell_borders(cell, TABLE_BORDER, "4")
            set_cell_margins(cell, 40, 40, 80, 80)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Set column widths proportionally
    total_width = Inches(6.3)
    if len(headers) <= 2:
        widths = [Inches(2.2), Inches(4.1)]
    elif len(headers) == 3:
        widths = [Inches(1.8), Inches(2.4), Inches(2.1)]
    elif len(headers) == 4:
        widths = [Inches(1.3), Inches(1.5), Inches(1.5), Inches(2.0)]
    elif len(headers) == 5:
        widths = [Inches(0.5), Inches(1.5), Inches(2.0), Inches(0.8), Inches(1.5)]
    elif len(headers) == 6:
        widths = [Inches(0.5), Inches(1.2), Inches(1.5), Inches(0.8), Inches(0.8), Inches(1.5)]
    else:
        w = total_width / len(headers)
        widths = [w] * len(headers)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                cell.width = widths[idx]

    doc.add_paragraph()  # spacer
    return table


def add_workflow_diagram(doc, title, states_str):
    """Create a visual workflow diagram with colored state boxes and arrows."""
    # Parse workflow: "State1 → State2 → State3"
    states = [s.strip() for s in re.split(r'→|->|➜', states_str) if s.strip()]
    if not states:
        return

    # Create a single-row table for the diagram
    num_cols = len(states) * 2 - 1  # states + arrows between them
    table = doc.add_table(rows=1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_idx = 0
    for i, state in enumerate(states):
        cell = table.rows[0].cells[col_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Get color for this state
        bg, fg = WORKFLOW_COLORS.get(state, ("E2E8F0", "334155"))

        run = p.add_run(state)
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(
            int(fg[0:2], 16), int(fg[2:4], 16), int(fg[4:6], 16)
        )
        run.font.name = 'Calibri'

        set_cell_shading(cell, bg)
        # Rounded-look borders matching the bg
        set_cell_borders(cell, bg, "8")
        set_cell_margins(cell, 30, 30, 40, 40)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        col_idx += 1

        # Arrow between states
        if i < len(states) - 1:
            arrow_cell = table.rows[0].cells[col_idx]
            arrow_cell.text = ""
            p2 = arrow_cell.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run("→")
            run2.font.size = Pt(14)
            run2.font.color.rgb = BRAND_PRIMARY
            run2.font.bold = True
            run2.font.name = 'Calibri'
            set_cell_borders(arrow_cell, "FFFFFF", "0")
            set_cell_margins(arrow_cell, 0, 0, 5, 5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            col_idx += 1

    doc.add_paragraph()  # spacer


def add_transition_table(doc, transitions):
    """Create a styled transition table for workflows."""
    if not transitions:
        return
    headers = list(transitions[0].keys())
    rows = [list(t.values()) for t in transitions]
    add_styled_table(doc, headers, rows, TABLE_HEADER_BG_ALT, compact=True)


def parse_inline(run_text, paragraph):
    """Parse inline markdown formatting (bold, code, italic) and add runs."""
    # Split by bold (**text**), code (`text`), italic (*text*)
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)', run_text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
            run.font.name = 'Calibri'
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
            run.font.name = 'Calibri'
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Calibri'


def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (headers, rows, end_idx)."""
    headers = []
    rows = []
    idx = start_idx

    # Header row
    if idx < len(lines) and '|' in lines[idx]:
        cells = [c.strip() for c in lines[idx].split('|')]
        cells = [c for c in cells if c]  # remove empty from leading/trailing |
        headers = cells
        idx += 1

    # Separator row (---|---|---)
    if idx < len(lines) and re.match(r'\s*\|?\s*[-:]+', lines[idx]):
        idx += 1

    # Data rows
    while idx < len(lines) and '|' in lines[idx]:
        line = lines[idx].strip()
        if not line or line.startswith('#'):
            break
        cells = [c.strip() for c in line.split('|')]
        cells = [c for c in cells if c != '']
        # Handle edge: leading/trailing pipes
        if line.startswith('|'):
            raw = line[1:]
            if raw.endswith('|'):
                raw = raw[:-1]
            cells = [c.strip() for c in raw.split('|')]
        rows.append(cells)
        idx += 1

    # Pad rows to match header length
    for i, row in enumerate(rows):
        while len(row) < len(headers):
            row.append("")
        rows[i] = row[:len(headers)]  # truncate if too many

    return headers, rows, idx


def parse_workflow_line(line):
    """Check if a line is a workflow diagram like: State1 → State2 → State3"""
    line = line.strip().strip('`')
    if '→' in line or '->' in line:
        parts = re.split(r'→|->|➜', line)
        parts = [p.strip() for p in parts if p.strip()]
        if len(parts) >= 2 and all(len(p) < 30 for p in parts):
            # Likely a workflow
            return True
    return False


def convert_markdown_to_docx(md_path, docx_path, title, subtitle):
    """Main conversion: parse markdown and build styled DOCX."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()
    setup_document(doc, title, subtitle)

    # Add Table of Contents placeholder
    toc_para = doc.add_paragraph()
    toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_para.add_run("TABLE OF CONTENTS")
    run.font.size = Pt(16)
    run.font.color.rgb = BRAND_PRIMARY
    run.font.bold = True
    run.font.name = 'Calibri'

    doc.add_paragraph()
    toc_note = doc.add_paragraph()
    toc_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_note.add_run("(Update this field in Word: Right-click → Update Field, or press Ctrl+A then F9)")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run.font.italic = True

    # Insert a TOC field
    p = doc.add_paragraph()
    run_elem = p.add_run()._r
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_elem.append(fldChar1)

    run_elem2 = p.add_run()._r
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run_elem2.append(instrText)

    run_elem3 = p.add_run()._r
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run_elem3.append(fldChar2)

    run_elem4 = p.add_run("(Table of contents - update in Word)")._r

    run_elem5 = p.add_run()._r
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_elem5.append(fldChar3)

    doc.add_page_break()

    idx = 0
    in_code_block = False
    code_buffer = []
    skip_toc_section = False

    while idx < len(lines):
        line = lines[idx]

        # Skip the markdown title line and metadata
        if idx < 5 and (line.startswith('# ') and 'Manual' in line or line.startswith('# ') and 'Guide' in line):
            idx += 1
            continue

        # Skip subtitle lines
        if line.startswith('**Enterprise') or line.startswith('**Version'):
            idx += 1
            continue

        # Skip markdown TOC
        if line.strip() == '## Table of Contents':
            skip_toc_section = True
            idx += 1
            continue

        if skip_toc_section:
            if line.strip().startswith('##') and 'Table of Contents' not in line:
                skip_toc_section = False
            elif line.strip() == '---':
                skip_toc_section = False
                idx += 1
                continue
            else:
                idx += 1
                continue

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_buffer)
                if code_text.strip():
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
                    # Add shading to the paragraph
                    pPr = p._p.get_or_add_pPr()
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9" w:val="clear"/>')
                    pPr.append(shading)
                    p.paragraph_format.left_indent = Cm(0.5)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
                code_buffer = []
            idx += 1
            continue

        if in_code_block:
            # Check if it's a workflow diagram inside code block
            stripped = line.strip()
            if parse_workflow_line(stripped):
                # End code block early, render as diagram
                if code_buffer:
                    code_text = '\n'.join(code_buffer)
                    if code_text.strip():
                        p = doc.add_paragraph()
                        run = p.add_run(code_text)
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                    code_buffer = []
                add_workflow_diagram(doc, "", stripped)
            else:
                code_buffer.append(line)
            idx += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            p = doc.add_paragraph()
            run = p.add_run("━" * 70)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
            idx += 1
            continue

        # Page break hint for major sections
        if line.startswith('## ') and not line.startswith('## Table'):
            # Add page break before major sections (## level)
            if idx > 10:  # Don't add before the very first section
                doc.add_page_break()

        # Headings
        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_heading(text, level=4)
            idx += 1
            continue
        elif line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            idx += 1
            continue
        elif line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            idx += 1
            continue
        elif line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            idx += 1
            continue

        # Tables
        if '|' in line and idx + 1 < len(lines) and ('---' in lines[idx + 1] or '|' in lines[idx + 1]):
            # Check if next line is separator
            next_line = lines[idx + 1] if idx + 1 < len(lines) else ""
            if re.match(r'\s*\|?\s*[-:| ]+\s*$', next_line):
                headers, rows, new_idx = parse_table(lines, idx)
                if headers and rows:
                    # Determine header color based on context
                    hdr_color = TABLE_HEADER_BG
                    if any(h.lower() in ['result', 'status', 'pass', 'fail'] for h in headers):
                        hdr_color = TABLE_HEADER_BG_ALT
                    compact = len(headers) > 4
                    add_styled_table(doc, headers, rows, hdr_color, compact)
                    idx = new_idx
                    continue

        # Blockquotes
        if line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            pPr = p._p.get_or_add_pPr()
            # Left border effect via indentation + color
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EFF6FF" w:val="clear"/>')
            pPr.append(shading)
            if text.startswith('⚠️') or text.startswith('**Caution'):
                shading2 = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FEF3C7" w:val="clear"/>')
                pPr.remove(shading)
                pPr.append(shading2)
            elif text.startswith('💡'):
                shading2 = parse_xml(f'<w:shd {nsdecls("w")} w:fill="ECFDF5" w:val="clear"/>')
                pPr.remove(shading)
                pPr.append(shading2)
            parse_inline(text, p)
            idx += 1
            continue

        # Bullet lists
        if line.strip().startswith('- '):
            text = line.strip()[2:].strip()
            # Check for checkbox
            if text.startswith('[ ] ') or text.startswith('[x] ') or text.startswith('[X] '):
                checked = text[1] in ('x', 'X')
                text = text[4:]
                marker = "☑ " if checked else "☐ "
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(marker)
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0x22, 0xC5, 0x5E) if checked else RGBColor(0x94, 0xA3, 0xB8)
                parse_inline(text, p)
            else:
                p = doc.add_paragraph(style='List Bullet')
                parse_inline(text, p)
            idx += 1
            continue

        # Numbered lists
        match = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
        if match:
            text = match.group(3)
            p = doc.add_paragraph(style='List Number')
            parse_inline(text, p)
            idx += 1
            continue

        # Inline workflow diagrams (not in code blocks)
        if parse_workflow_line(line) and not line.strip().startswith('|'):
            add_workflow_diagram(doc, "", line.strip())
            idx += 1
            continue

        # Empty lines
        if not line.strip():
            idx += 1
            continue

        # Regular paragraph
        if line.strip().startswith('*End of'):
            # End marker - style differently
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip().strip('*'))
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
            run.font.size = Pt(10)
            idx += 1
            continue

        p = doc.add_paragraph()
        parse_inline(line.strip(), p)
        idx += 1

    # Add footer info
    doc.add_paragraph()
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("━" * 50)
    run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
    run.font.size = Pt(8)

    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer2.add_run("EAM-CHI  •  Confidential  •  March 2026")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run.font.name = 'Calibri'

    doc.save(docx_path)
    print(f"✅ Created: {docx_path}")


def add_kpi_stat_row(doc, stats):
    """Add a row of KPI stat boxes (label + big number)."""
    num_cols = len(stats)
    table = doc.add_table(rows=2, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    STAT_COLORS = [
        ("1A56DB", "DBEAFE"),  # Blue
        ("0E7C86", "CCFBF1"),  # Teal
        ("7C3AED", "EDE9FE"),  # Purple
        ("D97706", "FEF3C7"),  # Amber
        ("059669", "D1FAE5"),  # Green
        ("DC2626", "FEE2E2"),  # Red
    ]

    for i, (value, label) in enumerate(stats):
        accent, bg = STAT_COLORS[i % len(STAT_COLORS)]

        # Value cell (big number)
        val_cell = table.rows[0].cells[i]
        val_cell.text = ""
        p = val_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(value))
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(int(accent[0:2], 16), int(accent[2:4], 16), int(accent[4:6], 16))
        run.font.name = 'Calibri'
        set_cell_shading(val_cell, bg)
        set_cell_borders(val_cell, accent, "6")
        set_cell_margins(val_cell, 60, 20, 40, 40)
        val_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Label cell
        lbl_cell = table.rows[1].cells[i]
        lbl_cell.text = ""
        p2 = lbl_cell.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(label)
        run2.font.size = Pt(8)
        run2.font.bold = True
        run2.font.color.rgb = RGBColor(int(accent[0:2], 16), int(accent[2:4], 16), int(accent[4:6], 16))
        run2.font.name = 'Calibri'
        set_cell_shading(lbl_cell, bg)
        set_cell_borders(lbl_cell, accent, "6")
        set_cell_margins(lbl_cell, 20, 60, 40, 40)
        lbl_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()


def add_scope_hierarchy_diagram(doc):
    """Add a visual RBAC scope hierarchy diagram: own → team → site → all."""
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    scopes = [
        ("own", "EF4444", "FEE2E2", "User's own\nrecords"),
        ("team", "F59E0B", "FEF3C7", "Department\nrecords"),
        ("site", "3B82F6", "DBEAFE", "Site-wide\nrecords"),
        ("all", "22C55E", "D1FAE5", "All records\nglobally"),
    ]

    col_idx = 0
    for i, (label, accent, bg, desc) in enumerate(scopes):
        cell = table.rows[0].cells[col_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run(label.upper())
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(int(accent[0:2], 16), int(accent[2:4], 16), int(accent[4:6], 16))
        run.font.name = 'Calibri'

        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(desc)
        run2.font.size = Pt(7.5)
        run2.font.color.rgb = BRAND_DARK
        run2.font.name = 'Calibri'

        set_cell_shading(cell, bg)
        set_cell_borders(cell, accent, "8")
        set_cell_margins(cell, 40, 40, 30, 30)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        col_idx += 1

        if i < len(scopes) - 1:
            arrow_cell = table.rows[0].cells[col_idx]
            arrow_cell.text = ""
            pa = arrow_cell.paragraphs[0]
            pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ra = pa.add_run("⊂")
            ra.font.size = Pt(16)
            ra.font.color.rgb = BRAND_PRIMARY
            ra.font.bold = True
            set_cell_borders(arrow_cell, "FFFFFF", "0")
            set_cell_margins(arrow_cell, 0, 0, 5, 5)
            col_idx += 1

    doc.add_paragraph()


def add_phase_banner(doc, phase_num, title, color_hex="1A56DB"):
    """Add a colored phase banner heading."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]

    run = p.add_run(f"  PHASE {phase_num}")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = 'Calibri'

    run2 = p.add_run(f"    {title}")
    run2.font.size = Pt(12)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run2.font.name = 'Calibri'

    set_cell_shading(cell, color_hex)
    set_cell_borders(cell, color_hex, "8")
    set_cell_margins(cell, 50, 50, 80, 80)

    doc.add_paragraph()


def add_security_badge(doc, round_num, color_hex, items):
    """Add a security round visual with colored badge and items."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Badge cell
    badge = table.rows[0].cells[0]
    badge.text = ""
    badge.width = Inches(1.0)
    p = badge.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Round\n{round_num}")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = 'Calibri'
    set_cell_shading(badge, color_hex)
    set_cell_borders(badge, color_hex, "8")
    set_cell_margins(badge, 40, 40, 30, 30)
    badge.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Items cell
    items_cell = table.rows[0].cells[1]
    items_cell.text = ""
    items_cell.width = Inches(5.3)
    for item in items:
        p2 = items_cell.add_paragraph() if items_cell.paragraphs[0].text else items_cell.paragraphs[0]
        run2 = p2.add_run(f"  ✓  {item}")
        run2.font.size = Pt(9)
        run2.font.name = 'Calibri'
        run2.font.color.rgb = BRAND_DARK
    light_bg = color_hex[0:2] + "15"  # Approximate light version
    set_cell_shading(items_cell, "F8FAFC")
    set_cell_borders(items_cell, color_hex, "4")
    set_cell_margins(items_cell, 30, 30, 60, 60)

    doc.add_paragraph()


def convert_features_report(md_path, docx_path):
    """Convert FEATURES_REPORT.md with enhanced visuals: KPI stats, scope diagrams, phase banners."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ── Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = BRAND_DARK
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15

    for level, (size, color) in {
        1: (22, BRAND_PRIMARY),
        2: (16, BRAND_SECONDARY),
        3: (13, BRAND_DARK),
        4: (11.5, RGBColor(0x37, 0x47, 0x51)),
    }.items():
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Calibri'
        h.font.size = Pt(size)
        h.font.color.rgb = color
        h.font.bold = True
        h.paragraph_format.space_before = Pt(24 if level == 1 else 18 if level == 2 else 14)
        h.paragraph_format.space_after = Pt(8 if level == 1 else 6)

    # ── Cover Page ──
    for _ in range(3):
        doc.add_paragraph()

    # Badge line
    badge_para = doc.add_paragraph()
    badge_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = badge_para.add_run("━━━  COMPOSABLE HOLDINGS INC.  ━━━")
    run.font.size = Pt(10)
    run.font.color.rgb = BRAND_SECONDARY
    run.font.name = 'Calibri'
    run.font.bold = True

    doc.add_paragraph()

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Added Features Report")
    run.font.size = Pt(36)
    run.font.color.rgb = BRAND_PRIMARY
    run.font.bold = True
    run.font.name = 'Calibri'

    # Subtitle
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run("Enterprise Asset Management — CHI & ITBA")
    run.font.size = Pt(16)
    run.font.color.rgb = BRAND_SECONDARY
    run.font.name = 'Calibri'

    doc.add_paragraph()

    # Version info
    ver_para = doc.add_paragraph()
    ver_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver_para.add_run("Version 1.0  •  March 24, 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    run.font.name = 'Calibri'

    scope_para = doc.add_paragraph()
    scope_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = scope_para.add_run("Commits: 49bde84 → 35ccd25  •  15 commits on main")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run.font.name = 'Calibri'

    # Decorative line
    doc.add_paragraph()
    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line_para.add_run("━" * 50)
    run.font.color.rgb = BRAND_PRIMARY
    run.font.size = Pt(12)

    doc.add_paragraph()

    # KPI stats on cover
    add_kpi_stat_row(doc, [
        ("22", "New\nEntities"),
        ("160+", "New\nColumns"),
        ("15", "RBAC\nRoles"),
        ("86", "Gaps\nResolved"),
        ("11", "New\nServices"),
        ("29.5K", "Lines\nAdded"),
    ])

    doc.add_page_break()

    # ── TOC ──
    toc_para = doc.add_paragraph()
    toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_para.add_run("TABLE OF CONTENTS")
    run.font.size = Pt(16)
    run.font.color.rgb = BRAND_PRIMARY
    run.font.bold = True
    run.font.name = 'Calibri'

    doc.add_paragraph()
    toc_note = doc.add_paragraph()
    toc_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_note.add_run("(Update this field in Word: Right-click → Update Field, or Ctrl+A then F9)")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run.font.italic = True

    p = doc.add_paragraph()
    run_elem = p.add_run()._r
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_elem.append(fldChar1)
    run_elem2 = p.add_run()._r
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run_elem2.append(instrText)
    run_elem3 = p.add_run()._r
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run_elem3.append(fldChar2)
    p.add_run("(Table of contents — update in Word)")
    run_elem5 = p.add_run()._r
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_elem5.append(fldChar3)

    doc.add_page_break()

    # ── Track special sections for enhanced rendering ──
    phase_colors = {
        "Phase 1": "1A56DB",
        "Phase 2": "0E7C86",
        "Phase 3+4": "7C3AED",
        "Phase 5": "D97706",
        "Phase 6": "059669",
    }
    rendered_scope_diagram = False

    idx = 0
    in_code_block = False
    code_buffer = []
    skip_header_line = True  # Skip the first # title since we have cover

    while idx < len(lines):
        line = lines[idx]

        # Skip the first markdown title and metadata lines
        if skip_header_line and (line.startswith('# EAM-CHI') or line.startswith('**Date:') or
                                  line.startswith('**Commits:') or line.startswith('**Scope:')):
            idx += 1
            continue
        if skip_header_line and line.strip() == '---':
            skip_header_line = False
            idx += 1
            continue

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_buffer)
                if code_text.strip():
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
                    pPr = p._p.get_or_add_pPr()
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9" w:val="clear"/>')
                    pPr.append(shading)
                    p.paragraph_format.left_indent = Cm(0.5)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
                code_buffer = []
            idx += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            idx += 1
            continue

        # Horizontal rule → thin decorative line
        if line.strip() == '---':
            p = doc.add_paragraph()
            run = p.add_run("━" * 70)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
            idx += 1
            continue

        # ── HEADINGS ──
        # H2 → page break + heading
        if line.startswith('## '):
            text = line[3:].strip()
            if idx > 10:
                doc.add_page_break()
            p = doc.add_heading(text, level=2)
            idx += 1
            continue

        # H3 with Phase banner detection
        if line.startswith('### '):
            text = line[4:].strip()
            # Check if it's a Phase heading
            phase_match = re.match(r'Phase\s+(\S+)\s*[—–-]\s*(.+)', text)
            if phase_match:
                phase_key = f"Phase {phase_match.group(1)}"
                phase_title = phase_match.group(2).strip()
                # Remove commit hash in parens
                phase_title = re.sub(r'\s*\(`[^`]+`\)\s*$', '', phase_title)
                color = phase_colors.get(phase_key, "1A56DB")
                add_phase_banner(doc, phase_match.group(1), phase_title, color)
            else:
                # Check for Security Round headings
                round_match = re.match(r'Round\s+(\d+)', text)
                if round_match:
                    doc.add_heading(text, level=3)
                else:
                    doc.add_heading(text, level=3)

            # Insert scope hierarchy diagram after "Scope Levels" heading
            if 'Scope Levels' in text and not rendered_scope_diagram:
                doc.add_paragraph()
                add_scope_hierarchy_diagram(doc)
                rendered_scope_diagram = True

            idx += 1
            continue

        if line.startswith('#### '):
            text = line[5:].strip()
            doc.add_heading(text, level=4)
            idx += 1
            continue

        # ── TABLES ──
        if '|' in line and idx + 1 < len(lines) and ('---' in lines[idx + 1] or '|' in lines[idx + 1]):
            next_line = lines[idx + 1] if idx + 1 < len(lines) else ""
            if re.match(r'\s*\|?\s*[-:| ]+\s*$', next_line):
                headers, rows, new_idx = parse_table(lines, idx)
                if headers and rows:
                    # Choose color based on table context
                    hdr_color = TABLE_HEADER_BG
                    if any(h.lower() in ['scope', 'visibility', 'use case'] for h in headers):
                        hdr_color = "059669"  # Green for RBAC tables
                    elif any(h.lower() in ['fix', 'category'] for h in headers):
                        hdr_color = "DC2626"  # Red for security/fix tables
                    elif any(h.lower() in ['service file', 'action handler'] for h in headers):
                        hdr_color = "7C3AED"  # Purple for service tables
                    elif any(h.lower() in ['dsl file', 'document', 'migration'] for h in headers):
                        hdr_color = TABLE_HEADER_BG_ALT  # Teal for reference tables
                    elif any(h.lower() in ['new entity', 'model'] for h in headers):
                        hdr_color = "D97706"  # Amber for entity tables
                    elif any(h.lower() in ['item'] for h in headers):
                        hdr_color = "0E7C86"  # Teal for infra
                    compact = len(headers) > 3
                    add_styled_table(doc, headers, rows, hdr_color, compact)
                    idx = new_idx
                    continue

        # ── BULLET LISTS ──
        if line.strip().startswith('- '):
            text = line.strip()[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            parse_inline(text, p)
            idx += 1
            continue

        # ── NUMBERED LISTS ──
        match = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
        if match:
            text = match.group(3)
            p = doc.add_paragraph(style='List Number')
            parse_inline(text, p)
            idx += 1
            continue

        # ── EMPTY LINES ──
        if not line.strip():
            idx += 1
            continue

        # ── REGULAR PARAGRAPH ──
        p = doc.add_paragraph()
        parse_inline(line.strip(), p)
        idx += 1

    # ── Footer ──
    doc.add_paragraph()
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("━" * 50)
    run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
    run.font.size = Pt(8)

    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer2.add_run("EAM-CHI  •  Features Report  •  Confidential  •  March 2026")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run.font.name = 'Calibri'

    doc.save(docx_path)
    print(f"✅ Created: {docx_path}")


def main():
    base_dir = "/Users/marlonceniza/Documents/Composable/EAM-CHI"

    # Convert Features Report
    fr_md = os.path.join(base_dir, "FEATURES_REPORT.md")
    fr_docx = os.path.join(base_dir, "FEATURES_REPORT.docx")
    if os.path.exists(fr_md):
        print("Converting Features Report...")
        convert_features_report(fr_md, fr_docx)
    else:
        print(f"❌ Not found: {fr_md}")

    # Convert User Manual
    um_md = os.path.join(base_dir, "USER_MANUAL.md")
    um_docx = os.path.join(base_dir, "USER_MANUAL.docx")
    if os.path.exists(um_md):
        print("Converting User Manual...")
        convert_markdown_to_docx(
            um_md, um_docx,
            "EAM-CHI User Manual",
            "Enterprise Asset Management System"
        )
    else:
        print(f"❌ Not found: {um_md}")

    # Convert Test Guide
    tg_md = os.path.join(base_dir, "TEST_GUIDE.md")
    tg_docx = os.path.join(base_dir, "TEST_GUIDE.docx")
    if os.path.exists(tg_md):
        print("Converting Test Guide...")
        convert_markdown_to_docx(
            tg_md, tg_docx,
            "EAM-CHI Test Guide",
            "System Validation & Workflow Testing"
        )
    else:
        print(f"❌ Not found: {tg_md}")

    # Convert Entity Relationships
    er_md = os.path.join(base_dir, "ENTITY_RELATIONSHIPS.md")
    er_docx = os.path.join(base_dir, "ENTITY_RELATIONSHIPS.docx")
    if os.path.exists(er_md):
        print("Converting Entity Relationships...")
        convert_markdown_to_docx(
            er_md, er_docx,
            "Entity Relationships",
            "Entity-to-Entity Data Model & Foreign Key Map"
        )
    else:
        print(f"❌ Not found: {er_md}")

    # Convert Workflow States
    ws_md = os.path.join(base_dir, "WORKFLOW_STATES.md")
    ws_docx = os.path.join(base_dir, "WORKFLOW_STATES.docx")
    if os.path.exists(ws_md):
        print("Converting Workflow States...")
        convert_markdown_to_docx(
            ws_md, ws_docx,
            "Workflow States & Entity Lifecycle",
            "State Machines, Transitions & Cross-Entity Cascades"
        )
    else:
        print(f"❌ Not found: {ws_md}")

    # Convert Gap Analysis Report
    ga_md = os.path.join(base_dir, "GAP_ANALYSIS_REPORT.md")
    ga_docx = os.path.join(base_dir, "GAP_ANALYSIS_REPORT.docx")
    if os.path.exists(ga_md):
        print("Converting Gap Analysis Report...")
        convert_markdown_to_docx(
            ga_md, ga_docx,
            "Gap Analysis Report",
            "DSL Specification Basis, 86 Gaps & Resolution Status"
        )
    else:
        print(f"❌ Not found: {ga_md}")

    print("\n🎉 Conversion complete!")


if __name__ == "__main__":
    main()
